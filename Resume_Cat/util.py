from docx import Document
import PyPDF2
import spacy
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def parse_word_document(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_pdf_document(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        for page in pdf.pages:
            text += page.extract_text()
    return text


nlp = spacy.load("en_core_web_sm")

def extract_keywords(text):
    return [token.lemma_ for token in nlp(text) if not token.is_stop and not token.is_punct]

def compare_keywords(resume_keywords, job_keywords):
    return set(resume_keywords).intersection(job_keywords)

def update_resume_text(resume_text, job_keywords):
    # Simplified example; real implementation needed
    return resume_text + "\nRelevant Skills: " + ", ".join(job_keywords)



def generate_word_document(updated_text, file_path):
    doc = Document()
    doc.add_paragraph(updated_text)
    doc.save(file_path)

def generate_pdf_document(updated_text, file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    text_obj = c.beginText(40, 750)
    text_obj.setFont("Helvetica", 10)
    for line in updated_text.split('\n'):
        text_obj.textLine(line)
    c.drawText(text_obj)
    c.save()
