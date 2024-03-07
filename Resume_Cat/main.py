from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
import spacy
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import shutil
import os
from uuid import uuid4

app = FastAPI()

UPLOAD_DIRECTORY = "temp_uploads"
if not os.path.exists(UPLOAD_DIRECTORY):
    os.makedirs(UPLOAD_DIRECTORY)

nlp = spacy.load("en_core_web_sm")

app.mount("/static", StaticFiles(directory="static"), name="static")

def validate_file(file: UploadFile):
    allowed_types = ["application/pdf", "application/msword", 
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Invalid file type")

@app.post("/upload/")
async def upload_file(resume: UploadFile = File(...), job_description: str = Form(...)):
    validate_file(resume)
    file_id = str(uuid4())
    file_location = f"{UPLOAD_DIRECTORY}/{file_id}"
    os.makedirs(file_location)
    with open(f"{file_location}/{resume.filename}", "wb") as buffer:
        shutil.copyfileobj(resume.file, buffer)
    with open(f"{file_location}/job_description.txt", "w") as jd_file:
        jd_file.write(job_description)
    return {"file_id": file_id, "detail": "Files uploaded successfully"}

# Assuming implementation details for parse_* and generate_* functions here
# Also assuming extract_keywords, compare_keywords, and update_resume_text implementations

@app.post("/process/")
async def process_resume(file_id: str):
    # This is a simplified placeholder for the actual processing logic
    return {"message": "Processing not implemented in this overview"}

@app.get("/download/{file_id}")
async def download_file(file_id: str):
    file_location = f"{UPLOAD_DIRECTORY}/{file_id}/updated_resume.pdf"
    if not os.path.exists(file_location):
        raise HTTPException(status_code=404, detail="Updated resume not found")
    return FileResponse(path=file_location, filename="updated_resume.pdf")

# +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++

def parse_word_document(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_pdf_document(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        for page in pdf.pages:
            text += page.extract_text()
    return text


nlp = spacy.load("en_core_web_sm")

def extract_keywords(text):
    return [token.lemma_ for token in nlp(text) if not token.is_stop and not token.is_punct]

def compare_keywords(resume_keywords, job_keywords):
    return set(resume_keywords).intersection(job_keywords)

def update_resume_text(resume_text, job_keywords):
    # Simplified example; real implementation needed
    return resume_text + "\nRelevant Skills: " + ", ".join(job_keywords)



def generate_word_document(updated_text, file_path):
    doc = Document()
    doc.add_paragraph(updated_text)
    doc.save(file_path)

def generate_pdf_document(updated_text, file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    text_obj = c.beginText(40, 750)
    text_obj.setFont("Helvetica", 10)
    for line in updated_text.split('\n'):
        text_obj.textLine(line)
    c.drawText(text_obj)
    c.save()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
